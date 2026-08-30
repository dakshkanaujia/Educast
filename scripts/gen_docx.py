#!/usr/bin/env python3
"""Generate an editable Word (.docx) version of the EduCast Final Project Report,
following the BITS Document-format template structure + formatting rules."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "docs/EduCast_Project_Report.docx"
doc = Document()

# ---------- Global formatting: Times New Roman 12, 1.5 spacing, 1" margins ----------
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = normal.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(6)

for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Inches(1)
    sec.left_margin = sec.right_margin = Inches(1)

# Restyle heading styles to Times New Roman, 14 / 13 pt
for name, size in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
    st = doc.styles[name]
    st.font.name = "Times New Roman"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor(0, 0, 0)
    st.paragraph_format.line_spacing = 1.5
    st.paragraph_format.space_before = Pt(10)
    st.paragraph_format.space_after = Pt(6)


def page_number_footer():
    """Center a PAGE field in the footer of the section."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def para(text, align=None, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def rich(segments):
    """segments: list of (text, bold, italic)."""
    p = doc.add_paragraph()
    for text, b, i in segments:
        r = p.add_run(text)
        r.bold = b
        r.italic = i
    return p


def h(text, level=1):
    return doc.add_heading(text, level=level)


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    return t


def toc_field():
    p = doc.add_paragraph()
    run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    hint = OxmlElement("w:t"); hint.text = "Right-click and choose 'Update Field' to build the Table of Contents."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for e in (b, instr, sep, hint, end):
        run._r.append(e)


page_number_footer()

# ================= COVER =================
para("EduCast", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=24)
para("A Demand-Driven, Real-Time Academic Support Marketplace",
     WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=14)
para("")
para("Capstone Project Report", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=15)
para("")
para("Submitted by:", WD_ALIGN_PARAGRAPH.CENTER, bold=True)
for nm in ["Daksh Kanaujia (Roll No. ____________)",
           "K Manoj Krishna (Roll No. ____________)",
           "Utkersh Basnet (2023EBCS010)"]:
    para(nm, WD_ALIGN_PARAGRAPH.CENTER)
para("")
for label, val in [("Program", "BSc Computer Science (Online Mode)"),
                   ("Institution", "BITS Pilani"),
                   ("Academic Year", "2025–2026"),
                   ("Internal Supervisor", "____________________")]:
    rich([(f"{label}: ", True, False), (val, False, False)]).alignment = WD_ALIGN_PARAGRAPH.CENTER
para("")
para("Date of Submission: 31 August 2026", WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ================= DECLARATION =================
h("Declaration", 1)
para('We hereby declare that this capstone project titled "EduCast — A Demand-Driven, '
     'Real-Time Academic Support Marketplace" is an original work carried out by us and '
     'has not been submitted to any other university or institution for the award of any degree.')
para("All external sources, libraries, and tools used have been duly acknowledged. Any "
     "AI-assisted or automated tooling used during development served as a productivity aid "
     "under our direction, and all resulting work was reviewed and tested by us.")
para("")
for nm, roll in [("Daksh Kanaujia", "____________"),
                 ("K Manoj Krishna", "____________"),
                 ("Utkersh Basnet", "2023EBCS010")]:
    para(f"Name: {nm}    Roll Number: {roll}")
    para("Signature: ____________________    Date: ____________")
    para("")
para("Countersigned (Internal Supervisor):")
para("Name: ____________________")
para("Signature: ____________________    Date: ____________")
doc.add_page_break()

# ================= ABSTRACT =================
h("Abstract", 1)
rich([("", False, False)])
para("Traditional e-learning platforms such as Udemy, Coursera, and YouTube are supply-driven: "
     "educators publish content speculatively and learners must hope the right material exists for "
     "their exact problem. This model fails the learner stuck on one specific, urgent thing who needs "
     "targeted help immediately. EduCast inverts this into a demand-driven, reverse-auction marketplace "
     "for academic help.")
rich([("Solution implemented: ", True, False),
      ("A student posts a specific problem as a bounty with a budget; mentors see it instantly and "
       "compete by placing bids; either side may negotiate via counter-offers; the student accepts the "
       "best bid, which atomically creates a private session room and holds a simulated escrow; after the "
       "session the student marks it complete and rates the mentor, releasing payment and publishing a "
       "review. Marketplace events propagate in real time over WebSockets.", False, False)])
rich([("Technologies used: ", True, False),
      ("Go (Gin) REST API with GORM over PostgreSQL, a Gorilla WebSocket hub, JWT auth with role-based "
       "access control, and a React Native (Expo) client, all containerised with Docker Compose.", False, False)])
rich([("Outcomes and results: ", True, False),
      ("The full loop — post, bid, negotiate, accept, collaborate, complete, rate — works end-to-end on "
       "the containerised stack. A 33-case black-box API validation suite passes at 100% (33/33).", False, False)])
doc.add_page_break()

# ================= TOC + LISTS =================
h("Table of Contents", 1)
toc_field()
doc.add_page_break()

h("List of Figures", 1)
add_table(["Figure", "Caption", "Section"],
          [["Fig 2.1", "High-level system architecture", "2.1"],
           ["Fig 2.2", "Data flow — bounty to completed session", "2.1"],
           ["Fig 2.3", "Application screenshots [insert]", "2.5"],
           ["Fig 4.1", "Deployment / demo screenshots [insert]", "4.0"],
           ["Fig 5.1", "GitHub commit history [insert]", "5.1"]],
          [1.0, 3.6, 1.2])
para("")
h("List of Tables", 1)
add_table(["Table", "Caption", "Section"],
          [["Table 2.1", "Technology stack", "2.2"],
           ["Table 2.2", "Database tables", "2.3"],
           ["Table 3.1", "Test cases and results", "3.2"],
           ["Table 5.1", "Weekly progress summary", "5.2"]],
          [1.0, 3.6, 1.2])
para("")
h("List of Abbreviations", 1)
add_table(["Abbr.", "Meaning"],
          [["API", "Application Programming Interface"], ["JWT", "JSON Web Token"],
           ["RBAC", "Role-Based Access Control"], ["ORM", "Object-Relational Mapping"],
           ["WS", "WebSocket"], ["CRUD", "Create, Read, Update, Delete"],
           ["UVP", "Unique Value Proposition"], ["VOD", "Video on Demand"]],
          [1.0, 4.8])
doc.add_page_break()

# ================= CHAPTER 1 =================
h("Chapter 1: Introduction", 1)
para("Note: Problem identification and initial system design were completed as part of the Study "
     "Project; this chapter summarises them and notes subsequent changes.", italic=True)
h("1.1 Overview of the Project", 2)
para("EduCast is a mobile-first, demand-driven marketplace for academic help. Instead of browsing a "
     "catalogue of pre-recorded courses, a learner posts a precise question as a bounty with a budget; "
     "qualified mentors compete by bidding; and the two sides move into a real-time, one-to-one session. "
     "It reframes online learning around the loop Ask → Compare → Negotiate → Accept → Learn.")
h("1.2 Problem Statement & Motivation", 2)
para("For learners, generic course libraries do not help when one is blocked on a single, specific thing, "
     "and there is no fast way to request targeted, live help. For mentors, producing content speculatively "
     "is high-effort and low-certainty; they want targeted, well-paid work matched to their expertise. For "
     "both, once a deal is struck they need trust and accountability: no double-booking, a tracked session, "
     "and a fair rating afterwards.")
h("1.3 Objectives of the Capstone", 2)
bullets([
    "Build a demand-driven marketplace where students post bounties and mentors bid.",
    "Propagate marketplace events in real time over WebSockets.",
    "Support price negotiation through counter-offers.",
    "Guarantee safe, atomic bid acceptance so a mentor cannot be double-booked.",
    "Simulate an escrow → release payment lifecycle.",
    "Persist ratings/reviews and expose mentor directories and profiles.",
    "Containerise the full stack for one-command startup.",
])
h("1.4 Scope of Implementation", 2)
rich([("In scope: ", True, False),
      ("authentication and role-based access, the full bounty/bid lifecycle, negotiation, atomic acceptance "
       "and session creation, in-session text chat, completion with rating and review, mentor directory and "
       "profiles, price-insight statistics, real-time events, and Docker deployment.", False, False)])
rich([("Out of scope (prototype): ", True, False),
      ("real payment processing, live video/audio streaming, push notifications, and full production "
       "hardening (rate limiting, secret management, HTTPS/WSS).", False, False)])
h("1.5 Organization of the Report", 2)
para("Chapter 2 details the implementation; Chapter 3 covers testing, validation and results; Chapter 4 "
     "describes execution and deployment; Chapter 5 presents project-execution evidence; Chapter 6 concludes "
     "with achievements, limitations and future work.")
doc.add_page_break()

# ================= CHAPTER 2 =================
h("Chapter 2: Implementation Details", 1)
h("2.1 System Architecture & Design", 2)
para("EduCast follows a three-tier architecture: a React Native (Expo) client communicates with a Go/Gin "
     "backend over REST (Axios) and a persistent WebSocket connection; the backend persists data in "
     "PostgreSQL through GORM. A single in-process WebSocket Hub tracks connected clients by user ID and role "
     "and delivers events to specific users (targeted) or to all clients (broadcast).")
para("[Insert high-level architecture diagram — Fig 2.1]", italic=True)
para("Data flow (Fig 2.2): Student posts a bounty → broadcast to online mentors → mentor bids → pushed to "
     "the student → optional counter-offers → student accepts (atomic: bounty→IN_PROGRESS, room created, "
     "ESCROW recorded, mentor notified) → session + chat → student completes + rates (bounty→CLOSED, RELEASE "
     "recorded, rating updated, review stored, both sides notified).")
h("2.2 Technology Stack", 2)
add_table(["Layer", "Technology", "Purpose"],
          [["Language (backend)", "Go 1.25", "High-concurrency API + WebSocket server"],
           ["Web framework", "Gin", "HTTP routing, middleware, JSON binding"],
           ["ORM / DB driver", "GORM + Postgres driver", "Data access layer"],
           ["Database", "PostgreSQL 16", "Relational persistence"],
           ["Real-time", "Gorilla WebSocket", "Live bounty/bid/session events"],
           ["Auth", "golang-jwt + bcrypt", "JWT sessions; password hashing"],
           ["Frontend", "React Native (Expo)", "Cross-platform client"],
           ["Infrastructure", "Docker & Docker Compose", "Containerised, one-command stack"]],
          [1.4, 1.9, 2.5])
para("Table 2.1 — Technology stack.", WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=9.5)
h("2.3 System Modules", 2)
bullets([
    "Authentication: signup/login, bcrypt hashing, JWT issuance and verification.",
    "Bounty: create/list/detail; role-scoped visibility; new bounties broadcast to mentors.",
    "Bid: mentors place bids (price, note, optional duration/preferred time); pushed to the student.",
    "Negotiation: either party counters; the other accepts/declines; no countering one's own pending counter.",
    "Acceptance / Session: atomic acceptance creates a room, records ESCROW, notifies the mentor.",
    "Completion / Review: student completes + rates; records RELEASE, updates the average, stores a review.",
    "Mentor directory / profile: real computed stats; search and subject filters.",
    "Real-time Hub: targeted and broadcast delivery, live presence, online-user counts.",
])
add_table(["Table", "Key columns"],
          [["users", "id, name, email (unique), password_hash, role, rating_avg, timestamps"],
           ["bounties", "id, student_id→users, title, description, subject_tag, budget, status, room_id"],
           ["bids", "id, bounty_id, mentor_id, price_offer, note, duration_minutes, preferred_time, counter_*, is_accepted"],
           ["transactions", "id, bounty_id→bounties, amount, type (ESCROW|RELEASE|REFUND), created_at"],
           ["reviews", "id, bounty_id (unique), mentor_id, student_id, rating (1–5), comment, created_at"]],
          [1.1, 4.7])
para("Table 2.2 — Database tables.", WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=9.5)
h("2.4 Key Algorithms / Logic", 2)
para("Atomic bid acceptance (prevents double-booking): the operation runs inside a database transaction "
     "that re-validates the bounty is still OPEN and the bid is not already accepted before generating a "
     "room, marking the bid accepted, recording the ESCROW transaction, and notifying the mentor. "
     "Counter-offer negotiation is guarded by a shared check that verifies the caller belongs to the bid and "
     "that a party cannot counter its own pending counter.")
h("2.5 Screenshots / Code Snippets", 2)
para("[Insert application screenshots — signup, student feed, mentor feed with live bids, negotiation, "
     "session/chat, completion & rating. See Fig 2.3.]", italic=True)
doc.add_page_break()

# ================= CHAPTER 3 =================
h("Chapter 3: Testing, Validation & Results", 1)
h("3.1 Test Plan", 2)
para("Testing strategy: black-box functional and validation testing of the REST API against a running "
     "instance of the containerised stack. Each test issues a real HTTP request and compares the actual "
     "response to the expected outcome. Coverage spans authentication, authorisation/RBAC, the full "
     "bounty/bid lifecycle, negotiation, atomic acceptance, completion/review, and directory/profile/insight.")
para("Tools used: a reproducible shell suite (scripts/api_test.sh) using curl and python3, executed against "
     "the Docker Compose deployment.")
h("3.2 Test Cases", 2)
TESTS = [
 ("TC-01","Sign up as Student","Valid Student payload","201 Created + token","PASS"),
 ("TC-02","Sign up as Mentor","Valid Mentor payload","201 Created + token","PASS"),
 ("TC-03","Reject duplicate email","Existing email","409 Conflict","PASS"),
 ("TC-04","Reject invalid email","email='bad'","400 Bad Request","PASS"),
 ("TC-05","Reject short password","password='1'","400 Bad Request","PASS"),
 ("TC-06","Reject invalid role","role='Admin'","400 Bad Request","PASS"),
 ("TC-07","Reject wrong password","Wrong password","401 Unauthorized","PASS"),
 ("TC-08","Login returns JWT","Valid credentials","Non-empty token","PASS"),
 ("TC-09","Block missing token","No Authorization header","401 Unauthorized","PASS"),
 ("TC-10","Block invalid token","Bad bearer token","401 Unauthorized","PASS"),
 ("TC-11","Mentor cannot post bounty","POST bounty as Mentor","403 Forbidden","PASS"),
 ("TC-12","Reject incomplete bounty","{title} only","400 Bad Request","PASS"),
 ("TC-13","Reject non-positive budget","budget=0","400 Bad Request","PASS"),
 ("TC-14","Create valid bounty","Full valid payload","Created, valid id","PASS"),
 ("TC-15","Fetch bounty by id","GET /bounties/:id","200 OK","PASS"),
 ("TC-16","Fetch missing bounty","GET /bounties/99999999","404 Not Found","PASS"),
 ("TC-17","Student cannot bid","POST bid as Student","403 Forbidden","PASS"),
 ("TC-18","Mentor places bid","price/note/duration/time","Created, valid id","PASS"),
 ("TC-19","Non-owner cannot view bids","GET bids as Mentor","403 Forbidden","PASS"),
 ("TC-20","Owner views bids","GET bids as owner","200 OK","PASS"),
 ("TC-21","Student counters bid","{price:350}","200 OK","PASS"),
 ("TC-22","Cannot counter own counter","Counter again","400 Bad Request","PASS"),
 ("TC-23","Mentor accepts counter","POST counter/accept","200 OK","PASS"),
 ("TC-24","Accept bid creates room","POST /bids/:id/accept","Non-empty room_id","PASS"),
 ("TC-25","No bids on non-open bounty","POST bid after accept","400 Bad Request","PASS"),
 ("TC-26","Double-accept prevented","Accept same bid again","400 Bad Request","PASS"),
 ("TC-27","Reject invalid rating","{rating:9}","400 Bad Request","PASS"),
 ("TC-28","Complete + rate","{rating:5, comment}","200 OK","PASS"),
 ("TC-29","Cannot re-complete","Complete CLOSED bounty","400 Bad Request","PASS"),
 ("TC-30","Mentor directory","GET /mentors","200 OK","PASS"),
 ("TC-31","Mentor profile","GET /mentors/:id","200 OK","PASS"),
 ("TC-32","Price insight","GET /price-insight?subject=Math","200 OK","PASS"),
 ("TC-33","Health check","GET /health","200 OK","PASS"),
]
add_table(["Test Case ID", "Description", "Input", "Expected Output", "Status"],
          [list(t) for t in TESTS], [0.9, 1.5, 1.4, 1.5, 0.7])
para("Table 3.1 — Test cases and results (33/33 passed).", WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=9.5)
h("3.3 Results & Analysis", 2)
para("Observations: all 33 test cases passed (100%). Validation rules (email format, password length, role "
     "enum, positive budget, rating 1–5) are enforced, and access-control boundaries hold. Atomic acceptance "
     "is confirmed indirectly by TC-25 and TC-26. Performance: API responses were observed in the "
     "sub-millisecond to low-millisecond range; real-time events were delivered to connected clients instantly.")
add_table(["Category", "Cases", "Passed", "Failed"],
          [["Authentication", 8, 8, 0], ["Authorisation / RBAC", 3, 3, 0],
           ["Bounty lifecycle", 5, 5, 0], ["Bidding", 4, 4, 0],
           ["Negotiation", 3, 3, 0], ["Acceptance / session", 3, 3, 0],
           ["Completion / review", 3, 3, 0], ["Directory / profile / insight", 4, 4, 0],
           ["Total", 33, 33, 0]], [2.6, 1.0, 1.0, 1.0])
doc.add_page_break()

# ================= CHAPTER 4 =================
h("Chapter 4: Execution / Deployment Details", 1)
h("4.1 Execution Environment", 2)
para("The system runs as three Docker containers orchestrated by Docker Compose: db (PostgreSQL 16, 5432), "
     "backend (Go/Gin, 8080), and frontend (Expo web, 8081). The backend waits for the database health check "
     "before starting; the database applies the schema from backend/migrations/ on first boot.")
h("4.2 Deployment Steps (Local)", 2)
bullets(["git clone https://github.com/dakshkanaujia/Educast.git",
         "cd Educast",
         "docker compose up --build",
         "curl http://localhost:8080/health   → {\"status\":\"ok\"}",
         "open http://localhost:8081 for the app"])
para("A native (non-Docker) path is also supported; see Appendix B / the Installation Guide, including the "
     "migration caveat for pre-existing database volumes.")
h("4.3 Demo Screenshots", 2)
para("[Insert deployment/demo screenshots — running containers, live feed, a real-time bid, an accepted "
     "session. See Fig 4.1.]", italic=True)
h("4.4 Demo Video Link", 2)
para("[Insert demo video URL here.]", italic=True)
doc.add_page_break()

# ================= CHAPTER 5 =================
h("Chapter 5: Project Execution Evidence", 1)
h("5.1 Version Control Evidence", 2)
rich([("GitHub repository: ", True, False), ("https://github.com/dakshkanaujia/Educast.git", False, False)])
para("[Insert a screenshot of the commit history. See Fig 5.1.]", italic=True)
h("5.2 Weekly Progress Summary", 2)
para("Fill in dates and obtain supervisor remarks. Entries reflect the actual milestones delivered.", italic=True)
add_table(["Week", "Task Planned", "Task Completed", "Supervisor Remark"],
          [["1", "Finalise design from Study Project", "Requirements & architecture confirmed", ""],
           ["2", "Backend scaffolding & auth", "Gin app, JWT auth, RBAC middleware", ""],
           ["3", "Bounty & bid lifecycle", "CRUD + role guards + validation", ""],
           ["4", "Real-time layer", "WebSocket hub; live bounty/bid events", ""],
           ["5", "Acceptance & escrow", "Atomic accept, room creation, ESCROW/RELEASE", ""],
           ["6", "Negotiation & reviews", "Counter-offers; ratings + mentor profiles", ""],
           ["7", "Dockerisation & DB", "Compose stack; MySQL→PostgreSQL migration", ""],
           ["8", "Testing & documentation", "33-case API suite (100%); full docs set", ""]],
          [0.55, 1.9, 2.15, 1.2])
para("Table 5.1 — Weekly progress summary.", WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=9.5)
h("5.3 Supervisor Interaction Summary", 2)
para("[Record review dates and key feedback received from the internal supervisor.]", italic=True)
add_table(["Review Date", "Key Feedback Received"],
          [["____________", ""], ["____________", ""], ["____________", ""]], [1.4, 4.4])
doc.add_page_break()

# ================= CHAPTER 6 =================
h("Chapter 6: Conclusion & Future Work", 1)
h("6.1 Summary of Implementation", 2)
para("EduCast delivers a working, real-time, demand-driven tutoring marketplace covering the complete loop "
     "from posting a bounty to a rated, completed session — with atomic acceptance, price negotiation, "
     "simulated escrow, and live mentor reputation — all containerised for one-command startup.")
h("6.2 Achievements", 2)
bullets([
    "Full marketplace loop implemented and working end-to-end.",
    "Real-time events over WebSockets, plus live presence.",
    "Atomic, race-safe bid acceptance with a simulated escrow → release lifecycle.",
    "Persisted ratings/reviews with mentor directory and profiles from real history.",
    "100% pass rate (33/33) on a reproducible API validation suite.",
    "Fully containerised deployment (PostgreSQL + backend + frontend).",
])
h("6.3 Limitations", 2)
bullets([
    "Payments and escrow are simulated as database rows (no real gateway).",
    "No live video/audio; the session room is a text/collaboration space.",
    "Real-time events reach only currently-connected clients (no persisted inbox).",
    "In-session chat is relayed but not persisted.",
    "A mentor may place multiple bids on the same bounty (no uniqueness guard).",
    "Production hardening pending (default secret, permissive CORS, no rate limiting, HTTP/WS).",
])
h("6.4 Future Enhancements", 2)
bullets([
    "Real payment gateway integration and a user wallet/balance.",
    "WebRTC live sessions using the existing signaling fields.",
    "Persistent notification inbox and chat history.",
    "Bounty upvoting, plus search and pagination on the feed.",
    "A migration runner (golang-migrate/goose) applied on startup.",
    "Automated Go unit/integration tests in CI.",
])
doc.add_page_break()

# ================= REFERENCES =================
h("References", 1)
for i, r in enumerate([
    "Gin Web Framework. https://github.com/gin-gonic/gin",
    "GORM — ORM library for Golang. https://gorm.io",
    "Gorilla WebSocket. https://github.com/gorilla/websocket",
    "golang-jwt. https://github.com/golang-jwt/jwt",
    "Expo — React Native framework. https://expo.dev",
    "PostgreSQL 16 Documentation. https://www.postgresql.org/docs/",
    "Docker Compose Documentation. https://docs.docker.com/compose/",
], 1):
    para(f"[{i}] {r}")
doc.add_page_break()

# ================= APPENDIX =================
h("Appendix", 1)
h("A. User Manual", 2)
para("A full user manual (Student and Mentor step-by-step flows) is provided with the deliverables "
     "(docs/05_User_Manual.md).")
h("B. Installation Guide", 2)
para("A full installation guide is provided (docs/06_Installation_Guide.md), covering the Docker and native "
     "paths, environment variables, the migration caveat for existing volumes, running the test suite, and "
     "troubleshooting.")
h("C. Source Code Link (GitHub)", 2)
para("https://github.com/dakshkanaujia/Educast.git")
h("D. Demo Video Link", 2)
para("[Insert demo video URL here.]", italic=True)

doc.save(OUT)
print("Saved", OUT)
